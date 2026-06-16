#!/usr/bin/env python3
"""Convert project BAO_CAO.md to Word (.docx)."""

from __future__ import annotations

import re
import sys
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

PROJECT_ROOT = Path(__file__).resolve().parents[1]
DEFAULT_INPUT = PROJECT_ROOT / "BAO_CAO.md"
DEFAULT_OUTPUT = PROJECT_ROOT / "BAO_CAO.docx"

TABLE_RE = re.compile(r"^\|(.+)\|$")
SEP_RE = re.compile(r"^\|[\s\-:|]+\|$")
CODE_FENCE_RE = re.compile(r"^```")


def set_doc_defaults(doc: Document) -> None:
    section = doc.sections[0]
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2)

    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(13)
    style.paragraph_format.line_spacing = 1.15
    style.paragraph_format.space_after = Pt(6)
    style._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")


def add_formatted_run(paragraph, text: str, bold: bool = False, italic: bool = False) -> None:
    """Parse **bold** and *italic* inline markers."""
    pattern = re.compile(r"(\*\*[^*]+\*\*|\*[^*]+\*|[^*]+)")
    for part in pattern.findall(text):
        if part.startswith("**") and part.endswith("**"):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        elif part.startswith("*") and part.endswith("*"):
            run = paragraph.add_run(part[1:-1])
            run.italic = True
        else:
            run = paragraph.add_run(part)
            if bold:
                run.bold = True
            if italic:
                run.italic = True


def parse_table_row(line: str) -> list[str]:
    cells = [c.strip() for c in line.strip().strip("|").split("|")]
    return cells


def add_table(doc: Document, rows: list[list[str]]) -> None:
    if not rows:
        return
    col_count = max(len(r) for r in rows)
    table = doc.add_table(rows=len(rows), cols=col_count)
    table.style = "Table Grid"
    for i, row in enumerate(rows):
        for j in range(col_count):
            text = row[j] if j < len(row) else ""
            cell = table.rows[i].cells[j]
            cell.text = text
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.name = "Times New Roman"
                    run.font.size = Pt(12)
                    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
                if i == 0:
                    for run in p.runs:
                        run.bold = True
    doc.add_paragraph()


def convert_md_to_docx(input_path: Path, output_path: Path) -> None:
    lines = input_path.read_text(encoding="utf-8").splitlines()
    doc = Document()
    set_doc_defaults(doc)

    in_code = False
    code_lines: list[str] = []
    table_rows: list[list[str]] | None = None
    in_html_div = False

    i = 0
    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        if stripped.startswith("<div"):
            in_html_div = True
            i += 1
            continue
        if stripped.startswith("</div"):
            in_html_div = False
            i += 1
            continue
        if in_html_div and stripped == "---":
            i += 1
            continue

        if CODE_FENCE_RE.match(stripped):
            if in_code:
                p = doc.add_paragraph()
                run = p.add_run("\n".join(code_lines))
                run.font.name = "Courier New"
                run.font.size = Pt(10)
                p.paragraph_format.left_indent = Cm(0.5)
                code_lines = []
                in_code = False
            else:
                in_code = True
            i += 1
            continue

        if in_code:
            code_lines.append(line)
            i += 1
            continue

        if TABLE_RE.match(stripped):
            if SEP_RE.match(stripped):
                i += 1
                continue
            if table_rows is None:
                table_rows = []
            table_rows.append(parse_table_row(stripped))
            i += 1
            continue

        if table_rows is not None:
            add_table(doc, table_rows)
            table_rows = None

        if stripped == "---":
            doc.add_paragraph()
            i += 1
            continue

        if stripped.startswith("# "):
            p = doc.add_heading(stripped[2:].strip(), level=0)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            i += 1
            continue

        if stripped.startswith("## "):
            doc.add_heading(stripped[3:].strip(), level=1)
            i += 1
            continue

        if stripped.startswith("### "):
            doc.add_heading(stripped[4:].strip(), level=2)
            i += 1
            continue

        if stripped.startswith("#### "):
            doc.add_heading(stripped[5:].strip(), level=3)
            i += 1
            continue

        if not stripped:
            i += 1
            continue

        if stripped.startswith("- ") or stripped.startswith("* "):
            p = doc.add_paragraph(style="List Bullet")
            add_formatted_run(p, stripped[2:])
            i += 1
            continue

        m = re.match(r"^(\d+)\.\s+(.+)$", stripped)
        if m:
            p = doc.add_paragraph(style="List Number")
            add_formatted_run(p, m.group(2))
            i += 1
            continue

        if stripped.startswith("```"):
            i += 1
            continue

        p = doc.add_paragraph()
        if in_html_div and i > 0:
            prev_nonempty = next(
                (lines[j].strip() for j in range(i - 1, -1, -1) if lines[j].strip()),
                "",
            )
            if prev_nonempty.startswith("#") or "**" in stripped:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER

        add_formatted_run(p, stripped)
        i += 1

    if table_rows is not None:
        add_table(doc, table_rows)

    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output_path))
    print(f"Đã tạo: {output_path}")


def main(argv: list[str] | None = None) -> int:
    args = argv if argv is not None else sys.argv[1:]
    input_path = Path(args[0]) if args else DEFAULT_INPUT
    output_path = Path(args[1]) if len(args) > 1 else DEFAULT_OUTPUT

    if not input_path.exists():
        print(f"Không tìm thấy: {input_path}", file=sys.stderr)
        return 1

    convert_md_to_docx(input_path, output_path)
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
